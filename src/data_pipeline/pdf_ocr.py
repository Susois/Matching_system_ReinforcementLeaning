"""
pdf_ocr.py
──────────
Extract text from PDF (and .docx) files in data/.
Strategy:
  1. Try pypdf (fast, works when PDF has selectable text).
  2. If extracted text is too short (< MIN_CHARS), fall back to
     pdf2image + pytesseract OCR (handles scanned PDFs).
  3. For .docx files use python-docx.

Outputs pure text strings — no Gemini calls here.
"""

import logging
import re
import warnings
from pathlib import Path
from typing import Optional

# Suppress pypdf "Multiple definitions in dictionary" noise
warnings.filterwarnings("ignore", message="Multiple definitions in dictionary")
logging.getLogger("pypdf").setLevel(logging.ERROR)

logger = logging.getLogger(__name__)

# Minimum characters to consider pypdf extraction successful
MIN_CHARS = 300


# ── pypdf (primary) ───────────────────────────────────────────
def _extract_with_pypdf(pdf_path: Path, max_pages: Optional[int] = None) -> str:
    """Extract text from a native/selectable PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(pdf_path))
        n = len(reader.pages)
        limit = min(max_pages, n) if max_pages else n
        parts = []
        for i in range(limit):
            try:
                text = reader.pages[i].extract_text() or ""
                parts.append(text)
            except Exception as e:
                logger.debug(f"pypdf page {i+1} error in {pdf_path.name}: {e}")
        return "\n".join(parts).strip()
    except ImportError:
        logger.warning("pypdf not installed; skipping native extraction.")
        return ""
    except Exception as e:
        logger.warning(f"pypdf failed for {pdf_path.name}: {e}")
        return ""


# ── pdf2image + pytesseract (fallback OCR) ───────────────────
def _extract_with_ocr(pdf_path: Path, max_pages: Optional[int] = None) -> str:
    """
    OCR fallback using pdf2image + pytesseract.
    Requires:  pip install pdf2image pytesseract
               Tesseract binary installed on system.
    Vietnamese language pack:  tessdata/vie.traineddata
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        logger.warning(
            "pdf2image / pytesseract not installed. "
            "Run: pip install pdf2image pytesseract  "
            "and install Tesseract + vie language pack for OCR fallback."
        )
        return ""

    try:
        images = convert_from_path(str(pdf_path), dpi=200)
        if max_pages:
            images = images[:max_pages]

        parts = []
        for idx, img in enumerate(images):
            # Try Vietnamese first, fallback to English
            for lang in ("vie+eng", "eng"):
                try:
                    text = pytesseract.image_to_string(img, lang=lang)
                    if text.strip():
                        parts.append(text)
                        break
                except pytesseract.TesseractError:
                    continue
            else:
                logger.debug(f"OCR produced no text for page {idx+1} of {pdf_path.name}")

        result = "\n".join(parts).strip()
        logger.info(f"OCR extracted {len(result)} chars from {pdf_path.name}")
        return result
    except Exception as e:
        logger.error(f"OCR failed for {pdf_path.name}: {e}")
        return ""


# ── python-docx ───────────────────────────────────────────────
def _extract_docx(docx_path: Path) -> str:
    """Extract text from a .docx file."""
    try:
        from docx import Document
        doc = Document(str(docx_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        logger.info(f"docx extracted {len(text)} chars from {docx_path.name}")
        return text
    except ImportError:
        logger.warning("python-docx not installed; cannot read .docx files.")
        return ""
    except Exception as e:
        logger.error(f"docx extraction failed for {docx_path.name}: {e}")
        return ""


# ── Public API ────────────────────────────────────────────────
def extract_text(file_path: Path, max_pages: Optional[int] = None) -> str:
    """
    Extract text from a PDF or DOCX file.
    Returns empty string on complete failure (never raises).
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx(file_path)

    if suffix != ".pdf":
        logger.warning(f"Unsupported file type: {file_path.name}")
        return ""

    # PDF: try native first
    text = _extract_with_pypdf(file_path, max_pages)

    if len(text) >= MIN_CHARS:
        logger.info(f"pypdf OK ({len(text)} chars): {file_path.name}")
        return text

    # Fallback to OCR
    logger.info(
        f"pypdf returned only {len(text)} chars for {file_path.name}. "
        "Trying OCR fallback..."
    )
    ocr_text = _extract_with_ocr(file_path, max_pages)
    if len(ocr_text) >= MIN_CHARS:
        return ocr_text

    # Return whichever is longer
    return ocr_text if len(ocr_text) > len(text) else text


def get_source_files(directory: Path, extensions: tuple = (".pdf", ".docx")) -> list[Path]:
    """
    Return sorted list of PDF/DOCX files from directory and its subdirectories.
    Scans both data/1.pdfs/ and data/2.docx/ if directory = data/.
    Skips zip, rar and other non-text formats.
    """
    files = []
    for f in sorted(directory.rglob("*")):
        if f.is_file() and f.suffix.lower() in extensions:
            # Skip files inside processed/ or embeddings/ subdirs
            if any(part in ("processed", "embeddings", "raw") for part in f.parts):
                continue
            files.append(f)
    logger.info(f"Found {len(files)} source files under {directory}")
    return files
